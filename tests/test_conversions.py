"""End-to-end conversions, asserted on the content that comes out."""

from __future__ import annotations

import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

from conftest import MARKER
from snapdox import Options, convert
from snapdox.errors import EncryptedPdf, ScannedPdf


def pdf_text(path: Path) -> str:
    import fitz

    with fitz.open(path) as doc:
        return " ".join(page.get_text("text") for page in doc)


def docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return " ".join(" ".join(parts).split())


# --- documents ---


def test_docx_to_pdf_keeps_the_text(fixtures, tmp_path):
    result = convert(fixtures / "sample.docx", tmp_path / "out.pdf")
    assert result.primary.exists()
    assert MARKER in " ".join(pdf_text(result.primary).split())


def test_pdf_to_docx_recovers_every_page(fixtures, tmp_path):
    result = convert(fixtures / "sample.pdf", tmp_path / "out.docx")
    text = docx_text(result.primary)
    for page in (1, 2, 3):
        assert f"Page {page} heading" in text


def test_pdf_to_docx_does_not_invent_tables(fixtures, tmp_path):
    """The default must not read ordinary columns as table cells."""
    from docx import Document

    result = convert(fixtures / "columns.pdf", tmp_path / "out.docx")
    assert len(Document(result.primary).tables) == 0


def test_pdf_to_docx_text_mode_reads_columns_in_order(fixtures, tmp_path):
    result = convert(
        fixtures / "columns.pdf", tmp_path / "out.docx", opts=Options(pdf_layout="text")
    )
    text = docx_text(result.primary)
    positions = [text.index(f"item {n}") for n in range(1, 13)]
    assert positions == sorted(positions), "columns were interleaved instead of read in order"


def test_xlsx_to_csv(fixtures, tmp_path):
    result = convert(fixtures / "sample.xlsx", tmp_path / "out.csv")
    assert "widget" in result.primary.read_text(encoding="utf-8")


def test_markdown_chains_to_pdf(fixtures, tmp_path):
    result = convert(fixtures / "sample.md", tmp_path / "out.pdf")
    assert not result.route.is_direct
    assert "Title" in pdf_text(result.primary)


# --- images ---


def test_pdf_to_png_writes_one_file_per_page(fixtures, tmp_path):
    result = convert(fixtures / "sample.pdf", tmp_path / "page.png", opts=Options(dpi=72))
    assert len(result.outputs) == 3
    assert [p.name for p in result.outputs] == ["page_p1.png", "page_p2.png", "page_p3.png"]


def test_page_selection_limits_the_output(fixtures, tmp_path):
    result = convert(fixtures / "sample.pdf", tmp_path / "p.png", opts=Options(pages="2", dpi=72))
    assert len(result.outputs) == 1
    assert result.primary.name == "p.png"  # single page keeps the plain name


def test_png_to_jpg_flattens_transparency(fixtures, tmp_path):
    from PIL import Image

    result = convert(fixtures / "sample.png", tmp_path / "out.jpg")
    with Image.open(result.primary) as img:
        assert img.mode == "RGB"


def test_png_to_svg_produces_real_paths(fixtures, tmp_path):
    result = convert(fixtures / "sample.png", tmp_path / "out.svg")
    root = ET.parse(result.primary).getroot()
    paths = root.findall(".//{http://www.w3.org/2000/svg}path")
    assert len(paths) >= 2, "expected one traced path per shape"
    assert all(p.get("d") for p in paths)


def test_svg_to_png_via_pdf(fixtures, tmp_path):
    from PIL import Image

    result = convert(fixtures / "sample.svg", tmp_path / "out.png", opts=Options(dpi=144))
    assert result.route.engines == ("svglib", "pymupdf")
    with Image.open(result.primary) as img:
        assert img.width > 200  # rendered above its 200pt natural width


def test_image_to_docx_embeds_the_picture(fixtures, tmp_path):
    from docx import Document

    result = convert(fixtures / "sample.png", tmp_path / "out.docx")
    assert len(Document(result.primary).inline_shapes) == 1


def test_png_to_ico(fixtures, tmp_path):
    result = convert(fixtures / "sample.png", tmp_path / "out.ico")
    assert result.primary.stat().st_size > 0


# --- failure paths ---


def test_scanned_pdf_is_refused_with_a_hint(fixtures, tmp_path):
    with pytest.raises(ScannedPdf) as exc:
        convert(fixtures / "scanned.pdf", tmp_path / "out.docx")
    assert "OCR" in exc.value.hint


def test_scanned_pdf_still_converts_to_images(fixtures, tmp_path):
    result = convert(fixtures / "scanned.pdf", tmp_path / "out.png", opts=Options(dpi=72))
    assert result.primary.exists()


def test_encrypted_pdf_needs_a_password(fixtures, tmp_path):
    with pytest.raises(EncryptedPdf):
        convert(fixtures / "locked.pdf", tmp_path / "out.docx")


def test_encrypted_pdf_opens_with_the_password(fixtures, tmp_path):
    result = convert(
        fixtures / "locked.pdf", tmp_path / "out.txt", opts=Options(password="secret")
    )
    assert MARKER.split()[1] in result.primary.read_text(encoding="utf-8")


def test_bad_page_range_is_reported(fixtures, tmp_path):
    with pytest.raises(ValueError, match="matches no pages"):
        convert(fixtures / "sample.pdf", tmp_path / "o.png", opts=Options(pages="99"))


def test_invalid_options_are_rejected_before_work_starts(fixtures, tmp_path):
    with pytest.raises(ValueError, match="dpi"):
        convert(fixtures / "sample.pdf", tmp_path / "o.png", opts=Options(dpi=99999))
