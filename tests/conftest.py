"""Test fixtures.

Sample files are generated rather than committed, so the suite carries no
binary blobs and every fixture's contents are visible right here.
"""

from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

#: Sentence planted in every text fixture so round-trips can be asserted.
MARKER = "The quick brown fox jumps over the lazy dog."


@pytest.fixture(scope="session")
def fixtures(tmp_path_factory) -> Path:
    out = tmp_path_factory.mktemp("fixtures")
    _make_docx(out / "sample.docx")
    _make_pdf(out / "sample.pdf")
    _make_two_column_pdf(out / "columns.pdf")
    _make_scanned_pdf(out / "scanned.pdf")
    _make_encrypted_pdf(out / "locked.pdf")
    _make_png(out / "sample.png")
    _make_xlsx(out / "sample.xlsx")
    _make_svg(out / "sample.svg")
    (out / "sample.md").write_text("# Title\n\nSome **bold** prose.\n", encoding="utf-8")
    return out


def _make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("SnapDox", level=1)
    doc.add_paragraph(MARKER)
    table = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            table.cell(r, c).text = f"r{r}c{c}"
    doc.save(path)


def _make_pdf(path: Path, pages: int = 3) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=letter)
    for page in range(pages):
        c.setFont("Helvetica-Bold", 18)
        c.drawString(72, 720, f"Page {page + 1} heading")
        c.setFont("Helvetica", 12)
        c.drawString(72, 690, MARKER)
        c.showPage()
    c.save()


def _make_two_column_pdf(path: Path) -> None:
    """Left column counts 1-6, right column 7-12, to prove reading order."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("Helvetica", 12)
    for i in range(6):
        c.drawString(60, 720 - i * 40, f"item {i + 1}")
    for i in range(6, 12):
        c.drawString(330, 720 - (i - 6) * 40, f"item {i + 1}")
    c.showPage()
    c.save()


def _make_scanned_pdf(path: Path) -> None:
    """A page containing only an image — no text layer at all."""
    import fitz
    from PIL import Image

    img_path = path.with_suffix(".png")
    Image.new("RGB", (600, 800), (240, 240, 240)).save(img_path)

    doc = fitz.open()
    page = doc.new_page(width=600, height=800)
    page.insert_image(fitz.Rect(0, 0, 600, 800), filename=str(img_path))
    doc.save(str(path))
    doc.close()
    img_path.unlink()


def _make_encrypted_pdf(path: Path) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 720), MARKER)
    doc.save(str(path), encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="secret")
    doc.close()


def _make_png(path: Path) -> None:
    from PIL import Image, ImageDraw

    img = Image.new("RGBA", (240, 160), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    draw.rectangle([20, 20, 120, 120], fill=(220, 40, 60, 255))
    draw.ellipse([130, 40, 220, 130], fill=(30, 90, 200, 255))
    img.save(path)


def _make_xlsx(path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["name", "qty"])
    ws.append(["widget", 4])
    wb.save(path)


def _make_svg(path: Path) -> None:
    path.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="200" height="120" '
        'viewBox="0 0 200 120"><rect x="10" y="10" width="80" height="80" fill="#dc283c"/>'
        '<circle cx="150" cy="60" r="45" fill="#1e5ac8"/></svg>',
        encoding="utf-8",
    )
