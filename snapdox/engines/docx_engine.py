"""PDF to Word, the headline conversion.

``pdf2docx`` reconstructs paragraphs, tables, images and column layout from the
PDF's own text and drawing operators.  That works beautifully on PDFs that were
*generated* from a document, and not at all on scans — there is simply no text
to recover.  Rather than emitting a Word file containing one giant picture of a
page, SnapDox detects that case up front and says so.
"""

from __future__ import annotations

import logging
import os
from contextlib import contextmanager
from pathlib import Path
from typing import Iterator

from ..errors import ConversionFailed, ScannedPdf
from ..formats import Kind
from ..options import Options, parse_pages
from ..registry import converter
from ._pdfutil import has_text_layer, open_pdf, ordered_text_blocks

_PDF2DOCX_MARKER = f"{os.sep}pdf2docx{os.sep}"


def _not_pdf2docx(record: logging.LogRecord) -> bool:
    return _PDF2DOCX_MARKER not in record.pathname


@contextmanager
def _quiet_pdf2docx() -> Iterator[None]:
    """Suppress pdf2docx's per-page narration.

    It calls ``logging.basicConfig`` at import and then logs through the root
    logger, so setting a level on its own logger achieves nothing.  Filtering
    by source file drops its chatter without touching anyone else's logging.
    """
    if os.environ.get("SNAPDOX_VERBOSE"):
        yield
        return

    handlers = list(logging.getLogger().handlers)
    for handler in handlers:
        handler.addFilter(_not_pdf2docx)
    try:
        yield
    finally:
        for handler in handlers:
            handler.removeFilter(_not_pdf2docx)


@converter(
    "pdf",
    "docx",
    # Named for the job rather than the library: this edge has three
    # strategies behind it and only two of them are pdf2docx.
    engine="pdf-to-word",
    weight=5,
    note="Rebuilds paragraphs, tables and images",
    needs_text=True,
)
def pdf_to_docx(src: Path, dst: Path, opts: Options) -> list[Path]:
    with open_pdf(src, opts.password) as doc:
        if doc.page_count == 0:
            raise ConversionFailed(f"{src.name} has no pages")
        if not has_text_layer(doc):
            raise ScannedPdf(f"{src.name} has no text layer to convert into Word")
        pages = parse_pages(opts.pages, doc.page_count)

        if opts.pdf_layout == "text":
            return _rebuild_as_text(doc, src, dst, pages)

        all_pages = pages == list(range(doc.page_count))

    return _rebuild_with_pdf2docx(src, dst, opts, pages, all_pages)


def _rebuild_with_pdf2docx(
    src: Path, dst: Path, opts: Options, pages: list[int], all_pages: bool
) -> list[Path]:
    """Reconstruct the page layout: paragraphs, images, tables, positioning."""
    from pdf2docx import Converter

    dst.parent.mkdir(parents=True, exist_ok=True)

    # Borderless-table detection is pdf2docx's biggest source of mangled
    # output: on an ordinary two-column document it reads the columns as table
    # cells, shreds paragraphs across them, and takes an order of magnitude
    # longer doing it. Only go looking for them when the user asks.
    settings = {"parse_stream_table": opts.pdf_layout == "tables"}

    with _quiet_pdf2docx():
        cv = Converter(str(src), password=opts.password)
        try:
            cv.convert(str(dst), pages=None if all_pages else pages, **settings)
        except Exception as exc:
            raise ConversionFailed(
                f"could not rebuild {src.name} as a Word document: {exc}",
                hint="Try --pdf-layout text for a plain, fully editable version.",
            ) from exc
        finally:
            cv.close()

    if not dst.is_file() or dst.stat().st_size == 0:
        raise ConversionFailed(f"pdf2docx wrote no output for {src.name}")

    return [dst]


#: Inside one block, lines this short are treated as list items rather than
#: wrapped prose, so answer options keep their own lines.
_LIST_LINE_CHARS = 40


def _block_paragraphs(block: str) -> list[str]:
    """Split one PDF text block into the paragraphs it should become.

    A block is usually a single wrapped paragraph, so its lines are joined.
    But a run of uniformly short lines is a list — answer options, a menu, an
    address — and joining those into one line would be wrong.
    """
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    if not lines:
        return []
    if len(lines) > 1 and all(len(line) <= _LIST_LINE_CHARS for line in lines):
        return lines

    joined = ""
    for line in lines:
        if not joined:
            joined = line
        elif joined.endswith("-"):
            joined = joined[:-1] + line  # rejoin a word split across lines
        else:
            joined += " " + line
    return [joined]


def _rebuild_as_text(doc, src: Path, dst: Path, pages: list[int]) -> list[Path]:
    """Write the PDF's text into a clean, unstyled, fully editable document.

    No frames, no tab-emulated columns, no absolute positioning — just
    paragraphs in reading order. Layout is gone, but nothing fights you when
    you edit it, which is the point.
    """
    from docx import Document

    document = Document()
    document.core_properties.title = src.stem

    for position, index in enumerate(pages):
        if position:
            document.add_page_break()
        for block in ordered_text_blocks(doc.load_page(index)):
            for paragraph in _block_paragraphs(block):
                document.add_paragraph(paragraph)

    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dst))
    return [dst]


#: Usable width of a Letter page with one-inch margins, in inches.
_PAGE_WIDTH_IN = 6.5

#: What Word assumes an image's pixels are worth when no DPI is recorded.
_ASSUMED_DPI = 96


@converter(Kind.RASTER, "docx", engine="python-docx", weight=5, note="Image placed on a Word page")
def image_to_docx(src: Path, dst: Path, opts: Options) -> list[Path]:
    """Put an image into a Word document, scaled to fit the page.

    Without this, ``png -> docx`` would route through the PDF hub into
    pdf2docx and fail every time, since a picture has no text to recover.
    """
    from docx import Document
    from docx.shared import Inches

    from . import raster

    image = raster.load(src)
    width_in = min(_PAGE_WIDTH_IN, image.width / _ASSUMED_DPI)

    document = Document()
    try:
        document.add_picture(str(src), width=Inches(width_in))
    except Exception:
        # python-docx only reads a handful of formats; hand it a PNG instead.
        import tempfile

        with tempfile.TemporaryDirectory(prefix="snapdox-docximg-") as tmp:
            staged = Path(tmp) / "image.png"
            raster.save(image, staged, opts)
            document.add_picture(str(staged), width=Inches(width_in))

    dst.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(dst))
    return [dst]
